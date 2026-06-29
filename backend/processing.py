"""
processing.py — File parsing & text extraction pipeline
Supports: PDF, DOCX, TXT, Images (OCR), Video/Audio (Whisper)
"""
import os
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ── PDF Extraction ─────────────────────────────────────────────────────
def extract_pdf(filepath: str) -> str:
    """Extract text from PDF using PyMuPDF with pytesseract fallback."""
    text = ""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        for page in doc:
            page_text = page.get_text("text")
            text += page_text + "\n"
        doc.close()
        if text.strip():
            logger.info(f"PDF extracted via PyMuPDF: {len(text)} chars")
            return text.strip()
    except Exception as e:
        logger.warning(f"PyMuPDF failed: {e}. Trying OCR fallback...")

    # Fallback: OCR each page as image
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io
        doc = fitz.open(filepath)
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text += pytesseract.image_to_string(img) + "\n"
        doc.close()
        logger.info(f"PDF extracted via OCR fallback: {len(text)} chars")
    except Exception as e:
        logger.error(f"PDF OCR fallback failed: {e}")
        raise RuntimeError(f"Could not extract text from PDF: {e}")

    return text.strip()


# ── DOCX Extraction ────────────────────────────────────────────────────
def extract_docx(filepath: str) -> str:
    """Extract text from DOCX files including tables."""
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts)
        logger.info(f"DOCX extracted: {len(text)} chars")
        return text
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise RuntimeError(f"Could not extract DOCX: {e}")


# ── TXT Extraction ─────────────────────────────────────────────────────
def extract_txt(filepath: str) -> str:
    """Read plain text file with encoding detection."""
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            with open(filepath, "r", encoding=encoding) as f:
                text = f.read()
            logger.info(f"TXT extracted ({encoding}): {len(text)} chars")
            return text
        except UnicodeDecodeError:
            continue
    raise RuntimeError("Could not decode text file with any supported encoding.")


# ── Tesseract Path Auto-Detection (Windows) ────────────────────────────
def _configure_tesseract() -> bool:
    """
    Try to locate and configure the Tesseract binary.
    Returns True if Tesseract is available, False otherwise.
    """
    import shutil
    import sys

    # Already in PATH?
    if shutil.which("tesseract"):
        return True

    # Common Windows install locations
    if sys.platform == "win32":
        import pytesseract
        candidates = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe".format(
                os.environ.get("USERNAME", "")
            ),
        ]
        for path in candidates:
            if os.path.isfile(path):
                pytesseract.pytesseract.tesseract_cmd = path
                logger.info(f"Tesseract found at: {path}")
                return True

    return False


def _ocr_with_easyocr(img) -> str:
    """Use easyocr as a fallback OCR engine (no system binary required)."""
    import easyocr
    import numpy as np

    reader = easyocr.Reader(["en"], gpu=False, verbose=False)
    result = reader.readtext(np.array(img), detail=0, paragraph=True)
    return "\n".join(result)


# ── Image OCR Extraction ───────────────────────────────────────────────
def extract_image(filepath: str) -> str:
    """
    Extract text from image.
    Strategy:
      1. Try pytesseract (auto-detect Tesseract binary on Windows paths)
      2. Fall back to easyocr (pure Python, no system binary needed)
    """
    from PIL import Image

    img = Image.open(filepath)
    # Upscale small images for better OCR accuracy
    w, h = img.size
    if w < 1000:
        img = img.resize((w * 2, h * 2), Image.LANCZOS)

    # ── Attempt 1: pytesseract ──
    try:
        import pytesseract
        if _configure_tesseract():
            text = pytesseract.image_to_string(img, config="--psm 6")
            logger.info(f"Image OCR via Tesseract: {len(text)} chars")
            return text.strip()
        else:
            logger.warning("Tesseract binary not found — trying easyocr fallback.")
    except Exception as e:
        logger.warning(f"Tesseract OCR failed ({e}) — trying easyocr fallback.")

    # ── Attempt 2: easyocr ──
    try:
        text = _ocr_with_easyocr(img)
        logger.info(f"Image OCR via easyocr: {len(text)} chars")
        return text.strip()
    except ImportError:
        raise RuntimeError(
            "No OCR engine available. Install either:\n"
            "  • Tesseract: https://github.com/UB-Mannheim/tesseract/wiki\n"
            "  • easyocr (pip install easyocr)"
        )
    except Exception as e:
        logger.error(f"easyocr failed: {e}")
        raise RuntimeError(f"Image OCR failed with all engines: {e}")


# ── Video/Audio Transcription ──────────────────────────────────────────
def extract_video(filepath: str) -> str:
    """Transcribe video/audio using OpenAI Whisper."""
    try:
        import whisper
        model = whisper.load_model("base")
        result = model.transcribe(filepath)
        text = result.get("text", "")
        logger.info(f"Whisper transcribed: {len(text)} chars")
        return text.strip()
    except ImportError:
        raise RuntimeError(
            "Whisper not installed. Run: pip install openai-whisper"
        )
    except Exception as e:
        logger.error(f"Whisper transcription failed: {e}")
        raise RuntimeError(f"Transcription failed: {e}")


# ── Language Detection ─────────────────────────────────────────────────
def detect_language(text: str) -> str:
    """Detect language of extracted text."""
    try:
        from langdetect import detect
        lang = detect(text[:2000])
        return lang
    except Exception:
        return "en"


# ── Master Dispatcher ──────────────────────────────────────────────────
def extract_text(filepath: str, file_type: str) -> dict:
    """
    Route file to correct extraction pipeline.
    Returns dict: {text, language, char_count, status}
    """
    extractors = {
        "pdf":   extract_pdf,
        "docx":  extract_docx,
        "txt":   extract_txt,
        "image": extract_image,
        "video": extract_video,
        "audio": extract_video,  # Whisper handles both
    }

    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    try:
        text = extractor(filepath)
        lang = detect_language(text)
        return {
            "text":       text,
            "language":   lang,
            "char_count": len(text),
            "status":     "success",
        }
    except Exception as e:
        logger.error(f"Extraction failed for {filepath}: {e}")
        return {
            "text":       "",
            "language":   "unknown",
            "char_count": 0,
            "status":     f"error: {str(e)}",
        }
